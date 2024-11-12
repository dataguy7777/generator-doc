import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import pandas as pd
import io
import os
from PIL import Image

# ---------------------------
# Session State Initialization
# ---------------------------
if 'cover_page' not in st.session_state:
    st.session_state['cover_page'] = None
if 'paragraphs' not in st.session_state:
    st.session_state['paragraphs'] = []
if 'tables' not in st.session_state:
    st.session_state['tables'] = []
if 'templates' not in st.session_state:
    st.session_state['templates'] = {}

# ---------------------------
# Helper Functions
# ---------------------------

def load_templates(folder_path):
    """
    Load .docx templates and corresponding images from the specified folder.
    Assumes that each .docx file has a corresponding image with the same base name.
    Supported image formats: .jpg, .jpeg, .png
    """
    templates = {}
    supported_image_formats = ['.jpg', '.jpeg', '.png']
    
    if not os.path.isdir(folder_path):
        st.error(f"The folder path '{folder_path}' does not exist or is not a directory.")
        return templates

    files = os.listdir(folder_path)
    docx_files = [f for f in files if f.lower().endswith('.docx')]

    for docx in docx_files:
        base_name = os.path.splitext(docx)[0]
        # Look for corresponding image
        image_path = None
        for ext in supported_image_formats:
            img_file = f"{base_name}{ext}"
            if img_file in files:
                image_path = os.path.join(folder_path, img_file)
                break
        if image_path:
            templates[base_name] = {
                "template_path": os.path.join(folder_path, docx),
                "image_path": image_path
            }
        else:
            st.warning(f"No corresponding image found for template '{docx}'. Skipping this template.")
    
    if not templates:
        st.error("No valid templates found in the specified folder.")
    return templates

def add_paragraph():
    st.subheader("Add Paragraph")
    parent_paragraph = st.text_area("Enter main paragraph:", key="parent_paragraph")
    if st.button("Add Paragraph"):
        if parent_paragraph.strip() != "":
            st.session_state['paragraphs'].append({
                "type": "paragraph",
                "content": parent_paragraph,
                "sub_paragraphs": [],
                "comments": []
            })
            st.success("Paragraph added!")
        else:
            st.error("Paragraph cannot be empty.")
    
    if st.session_state['paragraphs']:
        st.markdown("---")
        st.subheader("Existing Paragraphs")
        for idx, para in enumerate(st.session_state['paragraphs']):
            st.markdown(f"**Paragraph {idx + 1}:** {para['content']}")
            # Sub-paragraphs
            for sub_idx, sub in enumerate(para['sub_paragraphs']):
                st.markdown(f"&nbsp;&nbsp;&nbsp;**Sub-paragraph {idx + 1}.{sub_idx + 1}:** {sub}")
            # Comments
            for com_idx, com in enumerate(para['comments']):
                st.markdown(f"&nbsp;&nbsp;&nbsp;**Comment {idx + 1}.{com_idx + 1}:** {com}")

def add_sub_paragraph():
    st.subheader("Add Sub-Paragraph")
    if not st.session_state['paragraphs']:
        st.warning("Add a main paragraph first.")
    else:
        para_options = [f"Paragraph {i+1}" for i in range(len(st.session_state['paragraphs']))]
        para_idx = st.selectbox("Select Paragraph to Add Sub-Paragraph", para_options)
        sub_para = st.text_area("Enter sub-paragraph:", key="sub_paragraph")
        if st.button("Add Sub-Paragraph"):
            if sub_para.strip() != "":
                selected_index = para_options.index(para_idx)
                st.session_state['paragraphs'][selected_index]['sub_paragraphs'].append(sub_para)
                st.success("Sub-paragraph added!")
            else:
                st.error("Sub-paragraph cannot be empty.")

def add_comment():
    st.subheader("Add Comment")
    if not st.session_state['paragraphs']:
        st.warning("Add a main paragraph first.")
    else:
        para_options = [f"Paragraph {i+1}" for i in range(len(st.session_state['paragraphs']))]
        para_idx = st.selectbox("Select Paragraph to Add Comment", para_options)
        comment = st.text_area("Enter comment:", key="comment")
        if st.button("Add Comment"):
            if comment.strip() != "":
                selected_index = para_options.index(para_idx)
                st.session_state['paragraphs'][selected_index]['comments'].append(comment)
                st.success("Comment added!")
            else:
                st.error("Comment cannot be empty.")

def add_table():
    st.subheader("Add Table")
    with st.form("table_form"):
        num_rows = st.number_input("Number of rows", min_value=1, value=2, key="table_rows")
        num_cols = st.number_input("Number of columns", min_value=1, value=2, key="table_cols")
        table_data = []
        for row in range(int(num_rows)):
            row_data = []
            for col in range(int(num_cols)):
                cell = st.text_input(f"Row {row + 1}, Column {col + 1}", key=f"table_cell_{row}_{col}")
                row_data.append(cell)
            table_data.append(row_data)
        submitted = st.form_submit_button("Add Table")
        if submitted:
            df = pd.DataFrame(table_data, columns=[f"Column {i+1}" for i in range(int(num_cols))])
            st.session_state['tables'].append(df)
            st.success("Table added!")

def generate_preview():
    st.header("Preview Document")
    if st.session_state['cover_page']:
        st.subheader("Cover Page")
        st.image(st.session_state['cover_page']['image_path'], use_column_width=True)
    
    if st.session_state['paragraphs']:
        st.subheader("Paragraphs")
        for idx, para in enumerate(st.session_state['paragraphs']):
            st.markdown(f"**Paragraph {idx + 1}:** {para['content']}")
            for sub_idx, sub in enumerate(para['sub_paragraphs']):
                st.markdown(f"&nbsp;&nbsp;&nbsp;**Sub-paragraph {idx + 1}.{sub_idx + 1}:** {sub}")
            for com_idx, com in enumerate(para['comments']):
                st.markdown(f"&nbsp;&nbsp;&nbsp;**Comment {idx + 1}.{com_idx + 1}:** {com}")
    
    if st.session_state['tables']:
        st.subheader("Tables")
        for idx, table in enumerate(st.session_state['tables']):
            st.markdown(f"**Table {idx + 1}:**")
            st.table(table)

def create_word_document():
    if not st.session_state['cover_page']:
        doc = Document()
    else:
        doc = Document(st.session_state['cover_page']['template_path'])
    
    # Add paragraphs
    for para in st.session_state['paragraphs']:
        p = doc.add_paragraph(para['content'])
        p_format = p.paragraph_format
        p_format.space_after = Pt(12)
        for sub in para['sub_paragraphs']:
            sub_p = doc.add_paragraph(sub)
            sub_p.style = 'List Bullet'  # Example style for sub-paragraphs
        for com in para['comments']:
            comment_p = doc.add_paragraph(f"Comment: {com}")
            comment_p.italic = True  # Example styling for comments
    
    # Add tables
    for table in st.session_state['tables']:
        if isinstance(table, pd.DataFrame):
            table_doc = doc.add_table(rows=1, cols=len(table.columns))
            table_doc.style = 'Light List Accent 1'  # Example table style
            hdr_cells = table_doc.rows[0].cells
            for i, column in enumerate(table.columns):
                hdr_cells[i].text = str(column)
            for _, row in table.iterrows():
                row_cells = table_doc.add_row().cells
                for i, item in enumerate(row):
                    row_cells[i].text = str(item)
            doc.add_paragraph("")  # Add space after table
    
    # Save to a BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------------------
# Main Application
# ---------------------------

def main():
    st.title("Advanced Word Document Generator")

    # ---------------------------
    # Template Loading Section
    # ---------------------------
    st.sidebar.header("Cover Page Templates")
    folder_path = st.sidebar.text_input("Enter the path to your templates folder:", 
                                       value="", 
                                       help="Specify the local folder containing .docx templates and their preview images.")
    
    if folder_path:
        templates = load_templates(folder_path)
        if templates:
            st.session_state['templates'] = templates
            template_names = list(st.session_state['templates'].keys())
            cover_choice = st.sidebar.selectbox("Choose a Cover Page Template", template_names)
            st.session_state['cover_page'] = st.session_state['templates'][cover_choice]
            # Display preview
            st.sidebar.image(st.session_state['cover_page']['image_path'], use_column_width=True)
        else:
            st.session_state['templates'] = {}
            st.session_state['cover_page'] = None
    else:
        st.sidebar.info("Please enter the path to your templates folder to load available cover page templates.")

    st.sidebar.markdown("---")
    st.sidebar.header("Add Content")
    content_choice = st.sidebar.selectbox("Choose Content Type to Add", 
                                         ["Select", "Add Paragraph", "Add Sub-Paragraph", "Add Comment", "Add Table"])
    
    if content_choice == "Add Paragraph":
        add_paragraph()
    elif content_choice == "Add Sub-Paragraph":
        add_sub_paragraph()
    elif content_choice == "Add Comment":
        add_comment()
    elif content_choice == "Add Table":
        add_table()
    
    st.markdown("---")
    generate_preview()
    
    if st.button("Generate and Download Word Document"):
        if not (st.session_state['paragraphs'] or st.session_state['tables'] or st.session_state['cover_page']):
            st.error("No content to generate. Please add paragraphs, tables, or select a cover page.")
        else:
            doc_buffer = create_word_document()
            st.download_button(
                label="Download Word Document",
                data=doc_buffer,
                file_name="generated_document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
